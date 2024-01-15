from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re
from docx.shared import Pt,RGBColor
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from copy import deepcopy
from docx.enum.table import WD_TABLE_ALIGNMENT
import os, shutil

#提取名字
import openpyxl
wb = openpyxl.load_workbook('list.xlsx')
ws = wb.active
data=[]
for cell in ws['A']:
    data.append(str(cell.value))
#删除空格，将两个字中间插入两个空格
i=0
for name in data:
    data[i]=data[i].replace(" ","")
    if (len(data[i]) == 2):
        name_list=list(data[i])
        name_list.insert(1,'  ')
        data[i]=''.join(name_list)
    i=i+1
while len(data)% 50 !=0:
    data.append('')
document=Document(r"E:中文版席卡模板.docx")
document.styles['Normal'].font.size=Pt(150)
document.styles['Normal'].font.name= u'宋体'

def rewn(Name,re_num):
    tables=document.tables
    tables[re_num].cell(0,0).text=Name
    tables[re_num].cell(0,1).text=Name
    tables[re_num].cell(0,0).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
    tables[re_num].cell(0,1).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER
j=0
jt=0
q=0
for name in data:
    rewn(data[j],jt)
    j=j+1
    jt=jt+1
    if (j % 50 == 0):
        new_file_dir="E:\自动名牌编辑"
        new_name='out'+str(q)
        new_name=new_name+'.docx'
        new_file = os.path.join(new_file_dir, new_name)
        document.save(new_file)
        jt = jt - 50
        q = q + 1